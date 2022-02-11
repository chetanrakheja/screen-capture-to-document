#execute "pip install python-docx PyAutoGUI pynput autopy" if you face any issues in dependencies 
from docx import Document
from docx.shared import Inches
import pyautogui
import os
from pynput.keyboard import Key, Listener
import datetime

currentDir = os.getcwd() 
document = Document()

screenshotShotcut= Key.print_screen #Screnshot Key 
exit_combination_msg='ctrl_l(left ctrl) + Esc' #Exit Combination Message
exit_combination = {Key.ctrl_l, Key.esc} #Exit Combination [Left control and Esc combination]
defaultDocName = "testingScreenShots"
imgCount = 1 



# exit_combination = {Key.ctrl_l, Key.esc}
currently_pressed = set()


#on press function
def on_press(key):
	check_key(key)
	if key in exit_combination:
		currently_pressed.add(key)
		if currently_pressed == exit_combination:
			listener.stop()
			exitFun()
			
		
#on release function
def on_release(key):	
	try:
		currently_pressed.remove(key)
	except KeyError:
		pass 
		
# if key is ptrscn
def check_key(key):
	if key == screenshotShotcut:
		saveImg() 


# add screen shot to document
def addSSToDoc(imgPath):
	p = document.add_paragraph()#to add space between 2 screenshots
	r = p.add_run() 
	r.add_picture(imgPath,width=Inches(5.90551), height=Inches(3.54331)) #Change size of screenshot for view
	p = document.add_paragraph() #to add space between 2 screenshots

#save document
def saveDoc(fileName):
	document.save(MasterPath +"\\"+fileName+'.docx') # save Doc with fileName
	print()
	print("Document Saved at =>")
	print(MasterPath +"\\"+fileName+".docx ") #Printing Document Name

#on exit function
def exitFun():
	print("<===================================>")
	print("Press Return Key to save Document name as Default name ("+defaultDocName+")")
	fileName = input("Enter Document Name : ") # taking Document name
	bad_chars = [';', ':', '!', "*","<",">","\"","/","|","?","*"] #list of char which needs to be removed from doc file name
	for i in bad_chars : #Remove Bad Chars from document name
		fileName = fileName.replace(i, '')

	if(fileName==''): # if filename is not specified or got any error
		fileName = defaultDocName
	saveDoc(fileName)


def createDirs():
	global MasterPath
	now = datetime.datetime.now()
	currentDateTime = now.strftime("%d.%m.%y_%I%M%S%p")
	MasterPath = currentDir+"\\TestingData\\"+currentDateTime # create folder path of date time inside TestingData
	os.makedirs(MasterPath+"\\"+ "shots") #creating all the folder to avoid errors

# #save image as file and add image to doc
def saveImg():
	global imgCount
	shot = pyautogui.screenshot() #Take screenshot
	path = MasterPath +"\\shots\\" #update Path
	shot.save(path+'\\'+str(imgCount)+'.png') #Save Screenshot
	addSSToDoc(path+'//'+str(imgCount)+'.png') # add Screenhsot to Doc
	print('File Saved as ' +path+str(imgCount)+'.png') # Print Message of img saved
	imgCount=imgCount +1 #increase img count


def printStartMsg():
	print("<---------------------------------------------------->")
	print("-> Created by Chetan Rakheja | rakhejachetan@gmail.com | https://www.linkedin.com/in/chetanrakheja/")
	print("_______________________")
	print("Press PrtSc to take the Screenshot and save to folder and To Document")
	print("Press "+exit_combination_msg +" to exit and save the document")
	print("_______________________")
	print("Current Directory is =>"+os.getcwd())
	createDirs()
	

# Collect events until released
with Listener(on_press=on_press,on_release=on_release) as listener:
	printStartMsg()	# Print Personal Details Only Once
	listener.join()

